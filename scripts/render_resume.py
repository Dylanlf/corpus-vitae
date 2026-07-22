#!/usr/bin/env python3
"""
render_resume.py — turn a standard JSON Resume into an ATS-safe, single-column,
text-layer PDF and DOCX, per references/08-formatting.md.

Design: refined senior-technical spec — embedded clean grotesque (Noto Sans),
10pt body with generous leading, compact left-aligned header, near-monochrome,
0.5pt hairline rules under UPPERCASE section headings, asymmetric section spacing.
ATS-safety comes from STRUCTURE (single column, real text layer, standard headings)
+ a clean font ToUnicode map — verified with a copy-paste text test, not from a
"safe generic" look. Formatting only; never changes content; drops internal `meta`.

Usage:  python render_resume.py <resume.json> [--outdir DIR] [--no-credit]
Requires: python-docx, reportlab (see scripts/requirements.txt)
"""
import argparse, json, os, re

MONTHS = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
          "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
INK = "1A1A1A"     # near-black body (softer than pure black)
META = "555555"    # secondary meta (dates, title line)
RULE = "CCCCCC"    # hairline divider
CREDIT = "Drafted with corpus-vitae, an open-source résumé tool."

# candidate embeddable font families (first that fully exists wins)
FONT_SETS = [
    ("NotoSans", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
     "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
     "/usr/share/fonts/truetype/noto/NotoSans-Italic.ttf",
     "/usr/share/fonts/truetype/noto/NotoSans-BoldItalic.ttf"),
    ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
     "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf"),
]


def fmt_date(d):
    if not d:
        return ""
    m = re.match(r"^(\d{4})-(\d{2})", d)
    return f"{MONTHS[int(m.group(2))]} {m.group(1)}" if m else d


def date_range(item):
    s = fmt_date(item.get("startDate", ""))
    e = fmt_date(item.get("endDate", "")) or ("Present" if s else "")
    return " – ".join([x for x in (s, e) if x]) if s else e


def contact_line(basics):
    loc = basics.get("location", {}) or {}
    city = ", ".join(x for x in (loc.get("city"), loc.get("region")) if x)
    parts = [city, basics.get("phone", ""), basics.get("email", "")]
    for p in basics.get("profiles", []) or []:
        url = (p.get("url") or "").replace("https://", "").replace("http://", "").rstrip("/")
        if url:
            parts.append(url)
    if basics.get("url"):
        parts.append(basics["url"].replace("https://", "").replace("http://", ""))
    return "   ·   ".join(x for x in parts if x)


def skills_lines(skills):
    return [(g.get("name", ""), " · ".join(g.get("keywords", []) or [])) for g in (skills or [])]


# ----------------------------- PDF (reportlab) -----------------------------
def build_pdf(r, path, build_credit=True):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_RIGHT
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Table, TableStyle, HRFlowable)
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from xml.sax.saxutils import escape

    # register the first fully-present font family; fall back to Helvetica (base-14)
    fam, BOLD, ITAL = "Helvetica", "Helvetica-Bold", "Helvetica-Oblique"
    for name, reg, bold, ital, boldital in FONT_SETS:
        if all(os.path.exists(p) for p in (reg, bold, ital, boldital)):
            pdfmetrics.registerFont(TTFont(name, reg))
            pdfmetrics.registerFont(TTFont(name + "-Bd", bold))
            pdfmetrics.registerFont(TTFont(name + "-It", ital))
            pdfmetrics.registerFont(TTFont(name + "-BdIt", boldital))
            pdfmetrics.registerFontFamily(name, normal=name, bold=name + "-Bd",
                                          italic=name + "-It", boldItalic=name + "-BdIt")
            fam, BOLD, ITAL = name, name + "-Bd", name + "-It"
            break

    ink, meta, rule = HexColor("#" + INK), HexColor("#" + META), HexColor("#" + RULE)
    doc = SimpleDocTemplate(path, pagesize=letter,
                            topMargin=0.55 * inch, bottomMargin=0.55 * inch,
                            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
                            title=(r.get("basics", {}).get("name", "") + " — Résumé"))
    FULL = letter[0] - 1.4 * inch

    S = {
        "name":    ParagraphStyle("name", fontName=BOLD, fontSize=19, textColor=ink, leading=22, spaceAfter=2),
        "title":   ParagraphStyle("title", fontName=fam, fontSize=10, textColor=meta, leading=13, spaceAfter=2),
        "contact": ParagraphStyle("contact", fontName=fam, fontSize=9.5, textColor=ink, leading=12, spaceAfter=2),
        "head":    ParagraphStyle("head", fontName=BOLD, fontSize=10.5, textColor=ink, leading=12, spaceBefore=14, spaceAfter=3),
        "role":    ParagraphStyle("role", fontName=BOLD, fontSize=10, textColor=ink, leading=13),
        "dates":   ParagraphStyle("dates", fontName=fam, fontSize=9.5, textColor=meta, alignment=TA_RIGHT, leading=13),
        "orgsum":  ParagraphStyle("orgsum", fontName=ITAL, fontSize=9.5, textColor=meta, leading=12.5, spaceAfter=2),
        "body":    ParagraphStyle("body", fontName=fam, fontSize=10, textColor=ink, leading=13.7, spaceAfter=3),
        "bullet":  ParagraphStyle("bullet", fontName=fam, fontSize=10, textColor=ink, leftIndent=13, firstLineIndent=-9, leading=13.7, spaceAfter=3),
        "credit":  ParagraphStyle("credit", fontName=ITAL, fontSize=8, textColor=HexColor("#888888"), spaceBefore=12, leading=10),
    }
    E = escape
    flow = []

    def heading(t):
        flow.append(Paragraph(E(t.upper()), S["head"]))
        flow.append(HRFlowable(width="100%", thickness=0.5, color=rule, spaceBefore=0, spaceAfter=5))

    def role_row(title, dates, entry_gap=8):
        if dates:
            t = Table([[Paragraph(E(title), S["role"]), Paragraph(E(dates), S["dates"])]],
                      colWidths=[FULL - 1.7 * inch, 1.7 * inch])
            t.setStyle(TableStyle([("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                                   ("TOPPADDING", (0, 0), (-1, -1), entry_gap), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                                   ("VALIGN", (0, 0), (-1, -1), "BOTTOM")]))
            flow.append(t)
        else:
            p = ParagraphStyle("role_sp", parent=S["role"], spaceBefore=entry_gap)
            flow.append(Paragraph(E(title), p))

    b = r.get("basics", {})
    flow.append(Paragraph(E(b.get("name", "")), S["name"]))
    if b.get("label"):
        flow.append(Paragraph(E(b["label"]), S["title"]))
    flow.append(Paragraph(E(contact_line(b)), S["contact"]))

    if b.get("summary"):
        heading("Summary"); flow.append(Paragraph(E(b["summary"]), S["body"]))
    if r.get("skills"):
        heading("Skills")
        for nm, kws in skills_lines(r["skills"]):
            flow.append(Paragraph(f"<b>{E(nm)}:</b>  {E(kws)}", S["body"]))
    if r.get("work"):
        heading("Experience")
        for i, w in enumerate(r["work"]):
            title = "  —  ".join(x for x in (w.get("position"), w.get("name")) if x)
            role_row(title, date_range(w), entry_gap=(2 if i == 0 else 9))
            if w.get("summary"):
                flow.append(Paragraph(E(w["summary"]), S["orgsum"]))
            for h in w.get("highlights", []) or []:
                flow.append(Paragraph("•&nbsp;&nbsp;" + E(h), S["bullet"]))
    if r.get("projects"):
        heading("Projects")
        for pr in r["projects"]:
            url = (pr.get("url") or "").replace("https://", "").rstrip("/")
            role_row(pr.get("name", "") + (f"   ·   {url}" if url else ""), None, entry_gap=2)
            if pr.get("description"):
                flow.append(Paragraph(E(pr["description"]), S["body"]))
    if r.get("education"):
        heading("Education")
        for e in r["education"]:
            deg = " ".join(x for x in (e.get("studyType"), e.get("area")) if x)
            line = f"{deg}  —  {e.get('institution','')}".strip(" —")
            role_row(line, fmt_date(e.get("endDate", "")), entry_gap=2)
            for c in e.get("courses", []) or []:
                flow.append(Paragraph(E(c), S["orgsum"]))
    if r.get("certificates"):
        heading("Certifications")
        for c in r["certificates"]:
            txt = c.get("name", "") + (f" — {c['issuer']}" if c.get("issuer") else "") + \
                  (f" ({fmt_date(c.get('date',''))})" if c.get("date") else "")
            flow.append(Paragraph("•&nbsp;&nbsp;" + E(txt), S["bullet"]))
    if build_credit:
        flow.append(Paragraph(E(CREDIT), S["credit"]))
    doc.build(flow)


# ----------------------------- DOCX (python-docx) -----------------------------
def build_docx(r, path, build_credit=True, font="Arial"):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.55)
        s.left_margin = s.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15
    RIGHT_TAB = Inches(7.1)
    ink, meta = RGBColor.from_string(INK), RGBColor.from_string(META)

    def para(sa=3, sb=0):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(sb)
        return p

    def run(p, t, size=10, bold=False, italic=False, color=ink):
        rn = p.add_run(t); rn.font.size = Pt(size); rn.bold = bold; rn.italic = italic; rn.font.color.rgb = color
        return rn

    def hairline(p):
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "4"), ("w:space", "3"), ("w:color", RULE)):
            bottom.set(qn(k), v)
        pBdr.append(bottom); pPr.append(pBdr)

    def heading(t):
        p = para(sa=4, sb=12); run(p, t.upper(), size=10.5, bold=True); hairline(p)

    def role_row(title, dates, sb=9):
        p = para(sa=1, sb=sb)
        p.paragraph_format.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
        run(p, title, size=10, bold=True)
        if dates:
            run(p, "\t" + dates, size=9.5, color=meta)

    def bullet(t):
        p = para(sa=3); p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.first_line_indent = Inches(-0.13)
        run(p, "•  " + t, size=10)

    b = r.get("basics", {})
    pn = para(sa=2); run(pn, b.get("name", ""), size=19, bold=True)
    if b.get("label"):
        pl = para(sa=2); run(pl, b["label"], size=10, color=meta)
    pc = para(sa=4); run(pc, contact_line(b), size=9.5)

    if b.get("summary"):
        heading("Summary"); p = para(sa=4); run(p, b["summary"])
    if r.get("skills"):
        heading("Skills")
        for nm, kws in skills_lines(r["skills"]):
            p = para(sa=2); run(p, nm + ":  ", bold=True); run(p, kws)
    if r.get("work"):
        heading("Experience")
        for i, w in enumerate(r["work"]):
            role_row("  —  ".join(x for x in (w.get("position"), w.get("name")) if x),
                     date_range(w), sb=(2 if i == 0 else 9))
            if w.get("summary"):
                p = para(sa=2); run(p, w["summary"], italic=True, color=meta)
            for h in w.get("highlights", []) or []:
                bullet(h)
    if r.get("projects"):
        heading("Projects")
        for pr in r["projects"]:
            url = (pr.get("url") or "").replace("https://", "").rstrip("/")
            role_row(pr.get("name", "") + (f"   ·   {url}" if url else ""), "", sb=2)
            if pr.get("description"):
                p = para(sa=3); run(p, pr["description"])
    if r.get("education"):
        heading("Education")
        for e in r["education"]:
            deg = " ".join(x for x in (e.get("studyType"), e.get("area")) if x)
            role_row(f"{deg}  —  {e.get('institution','')}".strip(" —"), fmt_date(e.get("endDate", "")), sb=2)
            for c in e.get("courses", []) or []:
                p = para(sa=2); run(p, c, size=9.5, color=meta)
    if r.get("certificates"):
        heading("Certifications")
        for c in r["certificates"]:
            bullet(c.get("name", "") + (f" — {c['issuer']}" if c.get("issuer") else "") +
                   (f" ({fmt_date(c.get('date',''))})" if c.get("date") else ""))
    if build_credit:
        p = para(sb=12); run(p, CREDIT, size=8, italic=True, color=RGBColor.from_string("888888"))
    doc.save(path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("resume_json")
    ap.add_argument("--outdir", default=None)
    ap.add_argument("--no-credit", action="store_true")
    args = ap.parse_args()
    r = json.load(open(args.resume_json))
    r.pop("meta", None)  # never render internal build metadata (e.g. honestyDial)
    name = r.get("basics", {}).get("name", "Resume")
    slug = re.sub(r"\s+", "-", name.strip()) + "-Resume"
    outdir = args.outdir or os.path.dirname(os.path.abspath(args.resume_json))
    credit = not args.no_credit
    build_docx(r, os.path.join(outdir, slug + ".docx"), credit)
    build_pdf(r, os.path.join(outdir, slug + ".pdf"), credit)
    print("wrote:", os.path.join(outdir, slug + ".docx"))
    print("wrote:", os.path.join(outdir, slug + ".pdf"))


if __name__ == "__main__":
    main()
